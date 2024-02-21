import docx
from datetime import datetime 

def print_to_file(dates, events_list): # Convert the file type to docx look at this link https://stackoverflow.com/questions/47666642/adding-an-hyperlink-in-msword-by-using-python-docx
   doc = docx.Document()
   header = datetime.min
   format = "%A, %B %d, %Y"
   for event in events_list:
      if(dates.get_date_start() <= event.Dstart.date() and event.Dstart.date() <= dates.get_date_end()):
         if header.date() < event.Dstart.date():
            p =doc.add_paragraph() 
            p.add_run(f"{event.Dstart.strftime(format)}").font.bold = True
            header = event.Dstart
         if event.Name.lower() == "general meeting" or event.Name.lower() == "team meeting":
            p = doc.add_paragraph()
            add_hyperlink(p, event.Url, f"{event.Xhosts} {event.Name} at {str(event.Dstart.strftime('%I:%M %p')).lstrip('0')} in {event.Location}")
         else:
            p = doc.add_paragraph()
            add_hyperlink(p, event.Url, f"{event.Name} at {str(event.Dstart.strftime('%I:%M %p')).lstrip('0')} in {event.Location}")

   doc.save(f"This_Week_In_Student_Life_{dates.get_date_start()}.docx")

def add_hyperlink(paragraph, url, text):

   # This gets access to the document.xml.rels file and gets a new relation id value
   part = paragraph.part
   r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

   # Create the w:hyperlink tag and add needed values
   hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
   hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

   # Create a w:r element
   new_run = docx.oxml.shared.OxmlElement('w:r')

   # Create a new w:rPr element
   rPr = docx.oxml.shared.OxmlElement('w:rPr')

   # Join all the xml elements together add add the required text to the w:r element
   new_run.append(rPr)
   new_run.text = text
   hyperlink.append(new_run)

   paragraph._p.append(hyperlink)

   return hyperlink
